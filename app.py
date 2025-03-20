import streamlit as st
import os
import tempfile
import time
import pickle
import uuid
import base64
from io import BytesIO

# Library imports needed for audio processing
import soundfile
import pandas as pd
from pydub import AudioSegment
import torch
import numpy as np
import librosa
from scipy.signal import medfilt
import google.generativeai as genai
from dotenv import load_dotenv

import markdown
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

st.set_page_config(
    page_title="明沼沼翻譯機",
    page_icon="🎙️",
    layout="wide"
)

# Load environment variables from .env file
load_dotenv()

# Section: Sidebar Configuration
with st.sidebar:
    st.title("🎙️ 明沼沼翻譯機")
    st.markdown("好似唔明，但又好似聽得明的翻譯機")
    st.markdown("未來優化: ")

    # Dynamically add sidebar items
    sidebar_items = ["修復重複的摘要輸出", 
                     "Add Option to give feedback on transcribed_text for 嚟修正AI錯誤", 
                     "Add more input format"]
    for item in sidebar_items:
        st.markdown(f"- {item}")


# Get API keys
HF_API_KEY = os.environ.get("HF_API_KEY", "")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")

if not GEMINI_API_KEY:
    GEMINI_API_KEY = st.sidebar.text_input("Enter Gemini API Key:", type="password")
    if not GEMINI_API_KEY:
        st.sidebar.warning("⚠️ Please enter a valid Gemini API Key to use this application.")

# Initialize session state variables if they don't exist
if 'transcribed_text' not in st.session_state:
    st.session_state.transcribed_text = None
if 'refined_spoken_text' not in st.session_state:
    st.session_state.refined_spoken_text = None
if 'written_text' not in st.session_state:
    st.session_state.written_text = None
if 'refined_written_text' not in st.session_state:
    st.session_state.refined_written_text = None
if 'summary_text' not in st.session_state:
    st.session_state.summary_text = None
if 'temp_directory' not in st.session_state:
    st.session_state.temp_directory = tempfile.mkdtemp()
if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False
if 'current_filename' not in st.session_state:
    st.session_state.current_filename = None

# Functions from the notebook
def convert_audio_to_wav(input_file, output_file):
    """Convert .m4a file to .wav format."""
    audio = AudioSegment.from_file(input_file, format="m4a")
    audio.export(output_file, format="wav")
    return output_file

def get_all_file_paths(directory):
    file_paths = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith('.wav'):
                file_paths.append(os.path.join(root, file))
    file_paths.sort(key=lambda x: int(os.path.splitext(os.path.basename(x))[0].split('_')[-1]))
    return file_paths

def split_audio_by_silence(audio_path, output_dir, chunk_duration=60, min_silence_duration=0.5, 
                       silence_threshold=-40, padding=0.2):
    """
    Split an audio file into chunks based on silence detection.
    """
    os.makedirs(output_dir, exist_ok=True)
    
    # Load the audio file with pydub
    audio = AudioSegment.from_file(audio_path)
    
    # Convert to mono if stereo
    if audio.channels > 1:
        audio = audio.set_channels(1)
    
    # Get the duration in milliseconds
    total_duration_ms = len(audio)
    
    # Convert chunk duration to milliseconds
    chunk_duration_ms = chunk_duration * 1000
    
    # Load the audio file with librosa for silence detection
    y, sr = librosa.load(audio_path, sr=None)
    
    # Compute the RMS energy
    rms = librosa.feature.rms(y=y, frame_length=int(sr * 0.025), hop_length=int(sr * 0.01))[0]
    
    # Convert to dB
    rms_db = 20 * np.log10(rms + 1e-10)
    
    # Apply median filtering to smooth the energy curve
    rms_db_smoothed = medfilt(rms_db, kernel_size=11)
    
    # Find silence regions (below threshold)
    silence_mask = rms_db_smoothed < silence_threshold
    
    # Convert to time (seconds)
    times = librosa.times_like(rms_db_smoothed, sr=sr, hop_length=int(sr * 0.01))
    
    # Find continuous silence regions
    silence_regions = []
    in_silence = False
    start_time = 0
    
    for i, is_silent in enumerate(silence_mask):
        if is_silent and not in_silence:
            # Start of a silence region
            start_time = times[i]
            in_silence = True
        elif not is_silent and in_silence:
            # End of a silence region
            end_time = times[i]
            silence_duration = end_time - start_time
            if silence_duration >= min_silence_duration:
                silence_regions.append((start_time, end_time))
            in_silence = False
    
    # Add the last silence region if we're still in silence at the end
    if in_silence:
        end_time = times[-1]
        silence_duration = end_time - start_time
        if silence_duration >= min_silence_duration:
            silence_regions.append((start_time, end_time))
    
    # Sort silence regions by their start time
    silence_regions.sort(key=lambda x: x[0])
    
    # Determine the split points
    split_points = []
    current_chunk_end_ms = chunk_duration_ms
    
    # Find the closest silence region to each target chunk end
    for i in range(len(silence_regions)):
        silence_start_ms = silence_regions[i][0] * 1000
        silence_end_ms = silence_regions[i][1] * 1000
        
        # If this silence region is close to our target chunk end
        if abs(silence_start_ms - current_chunk_end_ms) < chunk_duration_ms * 0.2:
            split_points.append(silence_start_ms)
            current_chunk_end_ms += chunk_duration_ms
        
        # If we've already passed our target chunk end, use this silence region
        elif silence_start_ms > current_chunk_end_ms:
            split_points.append(silence_start_ms)
            current_chunk_end_ms = silence_start_ms + chunk_duration_ms
    
    # Add the start and end points
    split_points = [0] + split_points
    if split_points[-1] < total_duration_ms:
        split_points.append(total_duration_ms)
    
    # Generate chunks with progress bar
    chunk_paths = []
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i in range(len(split_points) - 1):
        # Calculate start and end times with padding
        start_ms = max(0, split_points[i] - padding * 1000)
        end_ms = min(total_duration_ms, split_points[i + 1] + padding * 1000)
        
        # Extract the chunk
        chunk = audio[start_ms:end_ms]
        
        # Generate the output filename
        output_filename = f"chunk_{i+1:03d}.wav"
        output_path = os.path.join(output_dir, output_filename)
        
        # Export the chunk
        chunk.export(output_path, format="wav")
        chunk_paths.append(output_path)
        
        # Update progress
        progress = (i + 1) / (len(split_points) - 1)
        progress_bar.progress(progress)
        status_text.text(f"Processing chunk {i+1}/{len(split_points)-1} - Duration: {(end_ms-start_ms)/1000:.2f}s")
        
    status_text.text("Audio chunking complete!")
    return chunk_paths

def transcribe_with_gemini(audio_chunk_paths, api_key):
    """Transcribe Cantonese audio using Gemini API"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-pro')
        
        prompt = """
        Transcribe this Cantonese audio in Hong Kong colloquial speech, removing filler words and repeated phrases. 
        Extract only spoken content and ignore background noise.
        """
        
        transcriptions = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, chunk_path in enumerate(audio_chunk_paths):
            # Upload audio file with specified MIME type
            f = genai.upload_file(chunk_path, mime_type='audio/wav')
            
            # Generate using multimodal capabilities
            status_text.text(f"Transcribing chunk {i+1}/{len(audio_chunk_paths)}...")
            response = model.generate_content([prompt, f])
            transcriptions.append(response.text)
            
            # Update progress
            progress = (i + 1) / len(audio_chunk_paths)
            progress_bar.progress(progress)
            
        status_text.text("Transcription complete!")
        return "".join(transcriptions)
    except Exception as e:
        st.error(f"Error with Gemini transcription: {str(e)}")
        return None

def gemini_prompt_call(message, api_key, prompt, temperature=0.2, top_p=0.7):
    """Generic function for Gemini API prompts"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-pro')
        
        prompt_parts = [prompt+'\n', message]
        
        generation_config = genai.GenerationConfig(
                            temperature=temperature,
                            top_p=top_p,
                            )  
        
        response = model.generate_content(prompt_parts, generation_config=generation_config)
        return response.text
        
    except Exception as e:
        st.error(f"An error occurred: {e}")
        return None

def refine_cantonese(message, api_key):
    """Refine Cantonese text"""
    prompt = """
            You are an expert in Cantonese. Please perform the following two steps and return the revised text.
            1. Correct potential inaccuracies in the Cantonese text in Hong Kong style. For example, in Hong Kong, text such as '噉', should be written as '咁'.
            2. Refine punctuation to the text.
            3. Remove repeated and unnecessary words.  For example, '係嘅係嘅係嘅' or '啱嘅啱嘅啱嘅' can be replaced with '係'.
            """
    return gemini_prompt_call(message, api_key, prompt)

def speech_to_written(message, api_key):
    """Convert spoken Cantonese to written Chinese"""
    prompt = """
            Convert the following spoken Cantonese text into standard written Chinese in traditional characters (zh-Hant). 
            Maintain grammatical correctness and coherence while ensuring the meaning is preserved. 
            Use formal written Chinese (書面語) rather than colloquial Cantonese (口語).
            """
    return gemini_prompt_call(message, api_key, prompt)

def refine_written_text(message, api_key):
    """Refine written Chinese text"""
    prompt = """
            Please refine the following Chinese text by making the sentence structure smoother and more coherent. 
            Reduce excessive punctuation, clarify ambiguous parts, and ensure natural readability while preserving the original meaning.
            Output in traditional Chinese characters (zh-Hant).
            """
    return gemini_prompt_call(message, api_key, prompt, temperature=1)

def summarize_refined_written(message, api_key, sum_lvl, bullet_points=False, output_format="md"):
    """Summarize text with customizable level of detail"""
    if output_format == "docx":
        if bullet_points:
            prompt = f"""Summarize the given Chinese text and output in Traditional Chinese (zh) with bullet points and sectioned format. 
            Use detail level {sum_lvl}, where 1 is the most concise and 5 is the most detailed. 
            Retain key information and context while ensuring clarity, coherence, and fluency."""
        else:
            prompt = f"""Summarize the given Chinese text and output in Traditional Chinese (zh) at summarization level {sum_lvl}, 
            where 1 is the most concise and 5 is the most detailed. 
            Retain key information and context while ensuring clarity, coherence, and fluency."""
    else:
        if bullet_points:
            prompt = f"""Summarize the given Chinese text and output in Traditional Chinese (zh) markdown format with bullet points and sectioned format. 
            Use detail level {sum_lvl}, where 1 is the most concise and 5 is the most detailed. 
            Retain key information and context while ensuring clarity, coherence, and fluency."""
        else:
            prompt = f"""Summarize the given Chinese text and output in Traditional Chinese (zh) markdown format at summarization level {sum_lvl}, 
            where 1 is the most concise and 5 is the most detailed. 
            Retain key information and context while ensuring clarity, coherence, and fluency."""

    return gemini_prompt_call(message, api_key, prompt)

def create_download_link(content, filename, format_type):
    """Create a download link for various file formats"""
    if format_type == "md":
        # For markdown, we can directly use the text
        b64 = base64.b64encode(content.encode()).decode()
        return f'<a href="data:text/markdown;base64,{b64}" download="{filename}.md">Download Markdown</a>'
    
    elif format_type == "txt":
        # For plain text
        b64 = base64.b64encode(content.encode()).decode()
        return f'<a href="data:text/plain;base64,{b64}" download="{filename}.txt">Download Text</a>'
    
    elif format_type == "docx":
        # Create a temporary markdown file
        with tempfile.NamedTemporaryFile(mode="w", delete=False, suffix=".md", encoding="utf-8") as tmp_md_file:
            tmp_md_file.write(content)
            tmp_md_path = tmp_md_file.name
        
        # Create a temporary docx file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx_file:
            tmp_docx_path = tmp_docx_file.name
        
        # Convert markdown to docx
        md_to_docx(tmp_md_path, tmp_docx_path)
        
        # Read the docx file as bytes
        with open(tmp_docx_path, "rb") as f:
            docx_bytes = f.read()
        
        # Encode to base64
        b64 = base64.b64encode(docx_bytes).decode()
        
        # Clean up temporary files
        os.remove(tmp_md_path)
        os.remove(tmp_docx_path)
        
        return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx">Download DOCX</a>'
    
    else:
        st.warning(f"Format {format_type} is not fully implemented in this demo version.")
        return None

import markdown
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

def md_to_docx(md_file, docx_file):
    """
    Converts a markdown file to a docx file, preserving bullet points and indentation.

    Args:
        md_file (str): Path to the input markdown file.
        docx_file (str): Path to the output docx file.
    """
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert markdown to HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])

        # Parse the HTML content using BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')

        # Create a new document
        document = Document()

        # Add the text to the document, preserving structure
        for element in soup.children:
            if element.name == 'h2':
                document.add_heading(element.text, level=0)
            elif element.name == 'h3':
                document.add_heading(element.text, level=1)
            elif element.name == 'ul' or element.name == 'ol':
                for li in element.find_all('li'):
                    document.add_paragraph(li.text, style='List Bullet')
            elif element.name == 'p':
                document.add_paragraph(element.text)

        # Save the document
        document.save(docx_file)

    except Exception as e:
        print(f"Error converting {md_file} to {docx_file}: {e}")       

# Section 1: File Upload
st.header("1️⃣ File Upload")
uploaded_file = st.file_uploader("Upload an M4A audio file", type=["m4a"])

if uploaded_file is not None:
    # Save the uploaded file to a temporary location
    temp_input_path = os.path.join(st.session_state.temp_directory, uploaded_file.name)
    with open(temp_input_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    # Store filename in session state for later use
    file_name = os.path.splitext(uploaded_file.name)[0]
    st.session_state.current_filename = file_name
    
    st.success(f"File '{uploaded_file.name}' successfully uploaded!")
    
    # Create conversion button
    if st.button("Process Audio"):
        # Convert to WAV
        cat_image = st.empty()
        with st.spinner("Converting audio format..."):
            cat_image.image("img/cat-0.gif", caption="Converting audio format...")
            temp_wav_path = os.path.join(st.session_state.temp_directory, f"{file_name}.wav")
            convert_audio_to_wav(temp_input_path, temp_wav_path)
        cat_image.empty()
        
        # Create a directory for chunks
        chunks_dir = os.path.join(st.session_state.temp_directory, f"chunks_{file_name}")
        os.makedirs(chunks_dir, exist_ok=True)
        
        # Split audio into chunks
        st.subheader("Splitting audio into chunks")
        splitting_cat_image = st.empty()
        with st.spinner("Splitting audio into chunks..."):
            splitting_cat_image.image("img/cat-1.gif", caption="Splitting audio into chunks...")
            chunk_paths = split_audio_by_silence(temp_wav_path, chunks_dir)
        splitting_cat_image.empty()

        # Transcription
        st.subheader("2️⃣ Transcription")
        
        # Transcribing
        cat_image = st.empty()
        with st.spinner("幫緊你 幫緊你..."):
            cat_image.image("img/cat.gif", caption="Transcribing audio...")
            st.session_state.transcribed_text = transcribe_with_gemini(chunk_paths, GEMINI_API_KEY)
        cat_image.empty()
        
        if st.session_state.transcribed_text:
            # Now we can move on without spinners
            with st.expander("View Raw Transcription", expanded=False):
                st.text_area("Raw Transcribed Text", st.session_state.transcribed_text, height=200)

            # Refine spoken text
            cat_image_2 = st.empty()
            cat_image_2.image("img/cat-2.gif", caption="Refining spoken Cantonese text...")
            st.session_state.refined_spoken_text = refine_cantonese(st.session_state.transcribed_text, GEMINI_API_KEY)
            cat_image_2.empty()

            with st.expander("View Refined Spoken Text", expanded=False):
                st.text_area("Refined Spoken Cantonese", st.session_state.refined_spoken_text, height=200)

            # Convert to written Chinese
            cat_image_3 = st.empty()
            cat_image_3.image("img/cat-3.gif", caption="Converting to written Chinese...")
            st.session_state.written_text = speech_to_written(st.session_state.refined_spoken_text, GEMINI_API_KEY)
            cat_image_3.empty()

            with st.expander("View Written Chinese", expanded=False):
                st.text_area("Written Chinese (First Pass)", st.session_state.written_text, height=200)

            # Refine written text
            cat_image_4 = st.empty()
            cat_image_4.image("img/cat-4.gif", caption="Refining written Chinese...")
            st.session_state.refined_written_text = refine_written_text(st.session_state.written_text, GEMINI_API_KEY)
            cat_image_4.empty()

            # Generate unique filename for written text
            written_text_filename = f"written_text_{uuid.uuid4()}.txt"
            st.session_state.written_text_path = os.path.join(st.session_state.temp_directory, written_text_filename)

            # Write refined written text to file
            with open(st.session_state.written_text_path, "w", encoding="utf-8") as f:
                f.write(st.session_state.refined_written_text)

            st.text_area("Final Written Chinese Text", st.session_state.refined_written_text, height=300)

            st.session_state.processing_complete = True
        else:
            st.error("Transcription failed. Please check your API key and try again.")

# Always show the transcription section if it's complete
if st.session_state.processing_complete:
    # Transcription section is already displayed above when processing completes
    # This ensures it persists even when moving to other sections
    with st.expander("View Raw Transcription", expanded=False):
        st.text_area("Raw Transcribed Text", st.session_state.transcribed_text, height=200, key="raw_transcribed_persist")

    with st.expander("View Refined Spoken Text", expanded=False):
        st.text_area("Refined Spoken Cantonese", st.session_state.refined_spoken_text, height=200, key="refined_spoken_persist")

    with st.expander("View Written Chinese", expanded=False):
        st.text_area("Written Chinese (First Pass)", st.session_state.written_text, height=200, key="written_chinese_persist")

    st.text_area("Final Written Chinese Text", st.session_state.refined_written_text, height=300, key="final_written_persist")

    # Section 3: Summarization
    st.header("3️⃣ Summarization")

    col1, col2 = st.columns(2)
    with col1:
        summary_format = st.radio("Summary Format", ["Paragraph Style", "Bullet Points"])
    with col2:
        detail_level = st.slider("Detail Level", min_value=1, max_value=5, value=3,
                               help="1 = Most concise, 5 = Most detailed")

    use_bullets = summary_format == "Bullet Points"

    if st.button("Generate Summary"):
        cat_image_6 = st.empty()
        with st.spinner("Generating summary..."):
            # Read from the saved file
            cat_image_6.image("img/cat-6.gif", caption="Generating summary...")
            if os.path.exists(st.session_state.written_text_path):
                with open(st.session_state.written_text_path, "r", encoding="utf-8") as f:
                    written_text_for_summary = f.read()
            else:
                st.error("Error: Written text file not found.")  # Should not happen, but good to have a check
                st.stop()

            st.session_state.summary_text = summarize_refined_written(
                written_text_for_summary,
                GEMINI_API_KEY,
                sum_lvl=detail_level,
                bullet_points=use_bullets,
                output_format=output_format if 'output_format' in locals() else "md"
            )
        cat_image_6.empty()


        st.markdown("### Summary")
        st.markdown(st.session_state.summary_text)

    # Section 4: Save Output
    if st.session_state.summary_text:
        st.header("4️⃣ Save Output")

        # Use the filename from session state
        default_filename = f"{st.session_state.current_filename}_summary" if st.session_state.current_filename else "audio_summary"

        col1, col2 = st.columns(2)
        with col1:
            output_format = st.selectbox("Output Format", ["md", "txt", "docx"])
        with col2:
            filename = st.text_input("Filename", value=default_filename)

        if st.button("Generate Download Link"):
            download_link = create_download_link(st.session_state.summary_text, filename, output_format)
            if download_link:
                st.markdown(download_link, unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown("🎙️ Cantonese Audio Transcriber | Created with Streamlit")
